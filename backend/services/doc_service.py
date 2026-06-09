"""
文档服务模块
负责文档的创建、解析、分块、向量化存储
支持 PDF/Word/TXT 格式文档处理
"""
import os
import asyncio
from typing import List, Optional
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, func, or_
from models.db_models import Document, DocumentChunk
from models.schemas import DocumentResponse, DocumentListResponse
from utils.splitter import split_text
from core.chroma_conn import add_documents, delete_documents
from core.config import get_settings
from core.logging_config import setup_logging

logger = setup_logging("doc_service")
settings = get_settings()


async def create_document(
    db: AsyncSession,
    tenant_id: int,
    title: str,
    file_name: str,
    file_path: str,
    file_size: int,
    file_type: str,
    user_id: int
) -> Document:
    """
    创建文档记录

    Args:
        db: 数据库会话
        tenant_id: 租户 ID
        title: 文档标题
        file_name: 原始文件名
        file_path: 文件存储路径
        file_size: 文件大小
        file_type: 文件类型
        user_id: 上传用户 ID

    Returns:
        创建的文档对象
    """
    doc = Document(
        tenant_id=tenant_id,
        title=title,
        file_name=file_name,
        file_path=file_path,
        file_size=file_size,
        file_type=file_type,
        status="pending",  # 初始状态为待处理
        created_by=user_id
    )
    db.add(doc)
    await db.flush()
    return doc


async def process_document(
    db: AsyncSession,
    document_id: int,
    text_content: str,
    tenant_id: int
) -> bool:
    """
    处理文档：分块 → 存储到数据库和向量库

    Args:
        db: 数据库会话
        document_id: 文档 ID
        text_content: 文档文本内容
        tenant_id: 租户 ID

    Returns:
        处理是否成功
    """
    logger.info(f"开始处理文档 document_id={document_id}, 文本长度={len(text_content)}")

    # 0. 空文本保护
    if not text_content or not text_content.strip():
        logger.warning(f"文档 {document_id} 文本内容为空，跳过处理")
        doc = (await db.execute(select(Document).where(Document.id == document_id))).scalar_one()
        doc.status = "failed"
        await db.flush()
        return False

    # 1. 计算内容哈希，检测重复
    import hashlib
    content_hash = hashlib.md5(text_content.strip().encode()).hexdigest()
    old_doc = (await db.execute(
        select(Document).where(
            Document.tenant_id == tenant_id,
            Document.content_hash == content_hash,
            Document.id != document_id,
        )
    )).scalar_one_or_none()
    if old_doc:
        logger.info(f"检测到重复文档：新 doc_id={document_id}，旧 doc_id={old_doc.id}，处理完成后清理旧数据")

    # 2. 更新当前文档的 content_hash
    current_doc = (await db.execute(select(Document).where(Document.id == document_id))).scalar_one()
    current_doc.content_hash = content_hash
    await db.flush()

    # 3. 删除该文档已有的 DB chunks（覆盖/重处理场景，防止唯一约束冲突）
    from sqlalchemy import delete as sql_delete
    await db.execute(
        sql_delete(DocumentChunk).where(DocumentChunk.document_id == document_id)
    )
    await db.flush()

    # 4. 文本分块
    chunks = split_text(text_content)
    logger.info(f"分块完成: {len(chunks)} 个 chunks")

    # 5. 先存储到向量库 (Chroma)，失败则直接返回
    try:
        metadatas = [
            {
                "document_id": str(document_id),
                "chunk_index": idx,
                "tenant_id": str(tenant_id)
            }
            for idx in range(len(chunks))
        ]
        logger.info(f"写入向量库: {len(chunks)} 条, tenant={tenant_id}")
        chunk_ids = [f"{document_id}_{idx}" for idx in range(len(chunks))]
        await asyncio.to_thread(
            add_documents,
            tenant_id=tenant_id,
            texts=chunks,
            metadatas=metadatas,
            ids=chunk_ids
        )
        logger.info(f"向量库写入成功")
    except Exception as e:
        err_msg = str(e) if str(e) else repr(e)
        logger.error(f"向量库存储失败: {type(e).__name__} - {err_msg}", exc_info=True)
        doc = (await db.execute(select(Document).where(Document.id == document_id))).scalar_one()
        doc.status = "failed"
        await db.flush()
        raise RuntimeError(f"向量库存储失败: {type(e).__name__} - {err_msg}")

    # 6. 向量存储成功后，存储分块到数据库
    for idx, chunk_text in enumerate(chunks):
        chunk = DocumentChunk(
            document_id=document_id,
            chunk_index=idx,
            text=chunk_text
        )
        db.add(chunk)

    # 7. 更新文档状态
    doc_result = await db.execute(select(Document).where(Document.id == document_id))
    doc = doc_result.scalar_one()
    doc.status = "completed"
    doc.chunk_count = len(chunks)
    await db.flush()

    # 8. 处理成功后，清理重复的旧文档
    if old_doc:
        logger.info(f"文档 {document_id} 处理完成，清理重复的旧文档 old_id={old_doc.id}")
        try:
            await asyncio.to_thread(
                delete_documents, tenant_id=tenant_id,
                where={"document_id": str(old_doc.id)}
            )
        except Exception as e:
            logger.warning(f"清理旧文档向量失败: {e}")
        from sqlalchemy import delete as sql_delete
        await db.execute(
            sql_delete(DocumentChunk).where(DocumentChunk.document_id == old_doc.id)
        )
        if old_doc.file_path and os.path.exists(old_doc.file_path):
            try:
                os.remove(old_doc.file_path)
            except OSError as e:
                logger.warning(f"删除旧文件失败: {e}")
        await db.delete(old_doc)
        await db.flush()
        from services.rag_service import invalidate_bm25_cache
        invalidate_bm25_cache(tenant_id)

    logger.info(f"文档 {document_id} 处理完成: {len(chunks)} 个 chunks")
    return True


async def get_document_by_id(db: AsyncSession, document_id: int, tenant_id: int) -> Optional[Document]:
    """获取文档详情（含全局共享文档）"""
    result = await db.execute(
        select(Document).where(
            Document.id == document_id,
            or_(Document.tenant_id == tenant_id, Document.tenant_id == 0)
        )
    )
    return result.scalar_one_or_none()


async def get_documents(
    db: AsyncSession,
    tenant_id: int,
    skip: int = 0,
    limit: int = 20,
    include_global: bool = True,
    only_global: bool = False
) -> DocumentListResponse:
    """获取文档列表（支持全局共享过滤）"""
    # 查询过滤
    if only_global:
        query_filter = Document.tenant_id == 0
    elif include_global:
        query_filter = or_(Document.tenant_id == tenant_id, Document.tenant_id == 0)
    else:
        query_filter = Document.tenant_id == tenant_id

    result = await db.execute(
        select(Document)
        .where(query_filter)
        .order_by(Document.created_at.desc())
        .offset(skip)
        .limit(limit)
    )
    documents = result.scalars().all()

    # 统计总数
    count_result = await db.execute(
        select(func.count()).select_from(Document).where(query_filter)
    )
    total = count_result.scalar() or 0

    return DocumentListResponse(
        total=total,
        items=[DocumentResponse.model_validate(d) for d in documents]
    )


async def delete_document(db: AsyncSession, document_id: int, tenant_id: int) -> bool:
    """删除文档及相关分块和向量"""
    doc = await get_document_by_id(db, document_id, tenant_id)
    if not doc:
        return False

    # 获取文件路径，用于后续删除物理文件
    file_path = doc.file_path

    # 删除分块记录
    chunks_result = await db.execute(
        select(DocumentChunk).where(DocumentChunk.document_id == document_id)
    )
    chunks = chunks_result.scalars().all()
    for chunk in chunks:
        await db.delete(chunk)

    # 删除向量库数据
    try:
        await asyncio.to_thread(
            delete_documents, tenant_id, where={"document_id": str(document_id)}
        )
    except Exception as e:
        logger.error(f"删除向量失败: {e}")

    # 删除物理文件
    if file_path and os.path.exists(file_path):
        try:
            os.remove(file_path)
        except Exception as e:
            logger.error(f"删除物理文件失败: {e}")

    # 删除文档记录
    await db.delete(doc)
    await db.flush()
    return True


def extract_text_from_file(file_path: str, file_type: str) -> str:
    """
    从文件中提取文本内容

    Args:
        file_path: 文件路径
        file_type: 文件类型 (pdf/docx/txt/jpg/png)

    Returns:
        提取的文本内容
    """
    text = ""
    logger.info(f"开始提取文本: file_path={file_path}, file_type={file_type}")
    try:
        if file_type == "pdf":
            from pypdf import PdfReader
            import pymupdf
            import tempfile
            from utils.ocr import OCRProcessor

            reader = PdfReader(file_path)
            pdf_doc = pymupdf.open(file_path)
            ocr = OCRProcessor()
            OCR_THRESHOLD = 50
            ocr_pages = 0

            try:
                for i, page in enumerate(reader.pages):
                    page_text = page.extract_text() or ""
                    if len(page_text.strip()) < OCR_THRESHOLD:
                        pix = pdf_doc[i].get_pixmap(dpi=200)
                        tmp_path = None
                        try:
                            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                                pix.save(tmp.name)
                                tmp_path = tmp.name
                            ocr_text = ocr.extract_text(tmp_path)
                        finally:
                            if tmp_path:
                                try:
                                    os.unlink(tmp_path)
                                except OSError:
                                    pass
                        if ocr_text.strip():
                            text += ocr_text + "\n"
                            ocr_pages += 1
                    else:
                        text += page_text + "\n"
            finally:
                pdf_doc.close()

            logger.info(f"PDF 提取完成: 总页数={len(reader.pages)}, OCR页数={ocr_pages}, 文本长度={len(text)}")
            return text.strip()

        elif file_type in ["docx", "doc"]:
            from docx import Document
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"

        elif file_type in ["txt", "md"]:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()

        elif file_type in ["jpg", "jpeg", "png", "bmp", "tiff"]:
            from utils.ocr import OCRProcessor
            ocr = OCRProcessor()
            text = ocr.extract_text(file_path)

    except Exception as e:
        logger.error(f"文本提取失败: {e}")
        raise

    logger.info(f"文本提取完成: {len(text)} 字符")
    return text